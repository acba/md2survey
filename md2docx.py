#!/usr/bin/env python3
"""
md2docx.py

Converte um arquivo SurveyMD em um documento .docx de revisão/impressão, com
formatação inspirada no anexo do questionário do TCE-RJ/SETIC.

Uso básico:
    python md2docx.py questionario_gerado.md questionario.docx

Uso com cabeçalho/logo extraído de um DOCX de referência:
    python md2docx.py questionario_gerado.md questionario.docx \
        --template-docx TSID03-ANEXO_QUESTIONARIO.docx

Observações:
- Este script gera um DOCX para leitura/revisão humana, não um formulário DOCX preenchível.
- O arquivo .lss continua sendo o artefato correto para importação no LimeSurvey.
- O script entende a macro [adoption] e expande a questão principal, justificativa
  de não aplicabilidade e detalhamento/checklist.

Dependências:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import html
import re
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from utils import (
    filter_survey_by_target as filter_target_survey,
    split_frontmatter,
    survey_targets,
    target_output_path,
    validate_target_config,
)


# -----------------------------------------------------------------------------
# Modelo intermediário
# -----------------------------------------------------------------------------

@dataclass
class Option:
    code: str
    text: str
    evidence: str = ""
    evidence_text: str = ""


@dataclass
class Scale:
    code: str
    type: str = "single"
    options: List[Option] = field(default_factory=list)


@dataclass
class Question:
    code: str
    type: str
    text_lines: List[str] = field(default_factory=list)
    help: str = ""
    mandatory: bool = False
    scale: str = ""
    visible_if: str = ""
    options: List[Option] = field(default_factory=list)
    subquestions: List[Option] = field(default_factory=list)
    detail_options: List[Option] = field(default_factory=list)
    subgroup: str = ""
    attrs: Dict[str, str] = field(default_factory=dict)


@dataclass
class Group:
    code: str
    title: str
    description_lines: List[str] = field(default_factory=list)
    questions: List[Question] = field(default_factory=list)


@dataclass
class Survey:
    meta: Dict[str, str] = field(default_factory=dict)
    title: str = ""
    scales: Dict[str, Scale] = field(default_factory=dict)
    groups: List[Group] = field(default_factory=list)


# -----------------------------------------------------------------------------
# Parser SurveyMD
# -----------------------------------------------------------------------------

def strip_comments(text: str) -> str:
    return re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)


def parse_option_line(line: str) -> Option:
    item = line.strip()
    if not item.startswith("-"):
        raise ValueError(f"Item de lista inválido: {line}")
    item = item[1:].strip()
    if "|" in item:
        code, text = item.split("|", 1)
    elif ":" in item:
        code, text = item.split(":", 1)
    else:
        raise ValueError(f"Use '- codigo | texto': {line}")
    return Option(clean_text(code.strip()), clean_text(text.strip()))


def clean_text(value: str) -> str:
    """Remove HTML simples e marcações Markdown básicas sem destruir acentos."""
    if value is None:
        return ""
    value = html.unescape(str(value))
    value = re.sub(r"<\s*br\s*/?>", "\n", value, flags=re.IGNORECASE)
    value = re.sub(r"</\s*p\s*>", "\n", value, flags=re.IGNORECASE)
    value = re.sub(r"<[^>]+>", "", value)
    value = value.replace("**", "")
    value = value.replace("__", "")
    value = value.replace(" ", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n\s+", "\n", value)
    return value.strip()


def parse_markdown(path: Path) -> Survey:
    text = strip_comments(path.read_text(encoding="utf-8"))
    meta, lines = split_frontmatter(text)
    survey = Survey(meta=meta, title=meta.get("title", "Questionário"))

    current_group: Optional[Group] = None
    current_question: Optional[Question] = None
    current_scale: Optional[Scale] = None
    mode: Optional[str] = None

    def finish_question():
        nonlocal current_question, current_group, mode
        if current_question is not None:
            if current_group is None:
                current_group = Group(code="g001", title="Questionário")
                survey.groups.append(current_group)
            current_group.questions.append(current_question)
        current_question = None
        mode = None

    def finish_scale():
        nonlocal current_scale, mode
        if current_scale is not None:
            survey.scales[current_scale.code] = current_scale
        current_scale = None
        mode = None

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            if current_question and mode is None:
                current_question.text_lines.append("")
            continue

        if stripped.startswith("# "):
            if not survey.title:
                survey.title = clean_text(stripped[2:])
            continue

        m_scale = re.match(r"^##\s+Escala:\s*([^\s]+)\s*$", stripped, flags=re.IGNORECASE)
        if m_scale:
            finish_question()
            finish_scale()
            current_scale = Scale(code=m_scale.group(1).strip())
            mode = "scale"
            continue

        m_group = re.match(r"^##\s+Grupo:\s*([^|]+?)(?:\s*\|\s*(.*))?$", stripped, flags=re.IGNORECASE)
        if m_group:
            finish_question()
            finish_scale()
            code = m_group.group(1).strip()
            title = clean_text(m_group.group(2) or code)
            current_group = Group(code=code, title=title)
            survey.groups.append(current_group)
            mode = None
            continue

        m_question = re.match(r"^###\s+([^\s]+)\s*\[([^\]]+)\]\s*$", stripped)
        if m_question:
            finish_question()
            finish_scale()
            current_question = Question(code=m_question.group(1).strip(), type=m_question.group(2).strip().lower())
            mode = None
            continue

        if current_scale is not None:
            if stripped.lower().startswith("type:"):
                current_scale.type = stripped.split(":", 1)[1].strip()
            elif stripped.startswith("-"):
                current_scale.options.append(parse_option_line(stripped))
            continue

        if current_group is not None and current_question is None and stripped.startswith(">"):
            current_group.description_lines.append(clean_text(stripped[1:].strip()))
            continue

        if current_question is not None:
            lower = stripped.lower()
            if lower in {"subquestions:", "detail_options:", "options:"}:
                mode = lower.rstrip(":")
                continue

            if mode in {"subquestions", "detail_options", "options"} and stripped.startswith("-"):
                opt = parse_option_line(stripped)
                if mode == "subquestions":
                    current_question.subquestions.append(opt)
                elif mode == "detail_options":
                    current_question.detail_options.append(opt)
                else:
                    current_question.options.append(opt)
                continue

            if ":" in stripped and not stripped.startswith("http://") and not stripped.startswith("https://"):
                key, value = stripped.split(":", 1)
                key = key.strip().lower()
                value = value.strip()
                known = {
                    "mandatory", "scale", "visible_if", "help", "evidence", "evidence_text", "evidence_if", "evidence_suffix",
                    "evidence_mandatory", "evidence_allowed_filetypes", "evidence_min_files", "evidence_max_files", "detail",
                    "detail_mandatory", "min_answers", "max_answers", "hide_tip", "allowed_filetypes",
                    "min_files", "max_files", "text", "question", "explain", "answer_width", "subgroup", "target",
                    "repeat_group_description",
                }
                if key in known:
                    if key == "help":
                        current_question.help = clean_text(value)
                    elif key == "subgroup":
                        current_question.subgroup = clean_text(value)
                    elif key == "mandatory":
                        current_question.mandatory = value.lower() in {"y", "yes", "true", "sim", "s", "1"}
                    elif key == "scale":
                        current_question.scale = value
                    elif key == "visible_if":
                        current_question.visible_if = value
                    elif key in {"text", "question"}:
                        current_question.text_lines.append(value)
                    elif key == "repeat_group_description":
                        pass
                    else:
                        current_question.attrs[key] = value
                    mode = None
                    continue

            current_question.text_lines.append(line)

    finish_question()
    finish_scale()
    validate_adoption_evidence_attrs(survey)
    validate_target_config(survey)
    return survey


def validate_adoption_evidence_attrs(survey: Survey) -> None:
    obsolete_evidence_keys = {
        "evidence", "evidence_mandatory", "evidence_allowed_filetypes",
        "evidence_min_files", "evidence_max_files", "evidence_suffix",
    }
    for group in survey.groups:
        for q in group.questions:
            if q.type != "adoption":
                continue
            if obsolete_evidence_keys.intersection(q.attrs):
                raise ValueError(
                    f"Campo evidence obsoleto em {q.code}: use evidence_text para descrever a evidência."
                )


# -----------------------------------------------------------------------------
# Geração DOCX
# -----------------------------------------------------------------------------

ADOPTION_OPTIONS = [
    Option("naoad", "Não adota"),
    Option("adfor", "Há decisão formal ou plano aprovado para adotá-lo"),
    Option("admen", "Adota em menor parte"),
    Option("adpar", "Adota parcialmente"),
    Option("admai", "Adota em maior parte ou totalmente"),
    Option("naoap", "Não se aplica"),
]

NSA_OPTIONS = [
    Option("A", "Não se aplica porque há lei e/ou norma, externa à organização, que impede a implementação desta prática."),
    Option("B", "Não se aplica porque há estudos que demonstram que o custo de implementar este controle é maior que o benefício que seria obtido dessa implementação."),
    Option("C", "Não se aplica por outras razões."),
]


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_paragraph_bottom_border(paragraph, color="000000", size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_font(run, *, bold=None, italic=None, size=None, color=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def resolve_style(doc: Document, *names: str) -> Optional[str]:
    for name in names:
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue
    return None


def add_text_paragraph(doc: Document, text: str = "", *, bold=False, italic=False, size=11, color="000000", align=None, space_after=3, space_before=0, left_indent=0, first_line_indent=None, style=None, name="Arial"):
    explicit_style = style is not None
    if style is None:
        style = resolve_style(doc, "List Paragraph")
    else:
        style = resolve_style(doc, style) or resolve_style(doc, "List Paragraph")
    p = doc.add_paragraph(style=style)
    if not explicit_style:
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_after = Pt(space_after)
        pf.space_before = Pt(space_before)
        pf.line_spacing = 1.08
        pf.left_indent = Cm(left_indent)
        if first_line_indent is not None:
            pf.first_line_indent = Cm(first_line_indent)
    for i, part in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(part)
        if not explicit_style:
            set_font(run, bold=bold, italic=italic, size=size, color=color, name=name)
    return p


def add_blank(doc, height=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(height)
    p.paragraph_format.space_before = Pt(0)
    return p


def start_question_page(doc: Document):
    doc.add_page_break()


def extract_logo_from_template(template_docx: Optional[Path]) -> Optional[Path]:
    if not template_docx or not template_docx.exists():
        return None
    tmp_dir = Path(tempfile.mkdtemp(prefix="md2docx_logo_"))
    with zipfile.ZipFile(template_docx) as z:
        names = [n for n in z.namelist() if n.startswith("word/media/") and n.lower().endswith((".png", ".jpg", ".jpeg"))]
        # No documento de referência, image2.png é o logotipo pequeno do cabeçalho.
        names_sorted = sorted(names, key=lambda n: (0 if "image2" in n.lower() else 1, n))
        for name in names_sorted:
            out = tmp_dir / Path(name).name
            out.write_bytes(z.read(name))
            return out
    return None


def configure_document(doc: Document, logo_path: Optional[Path]):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.bold = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(11)

    header = section.header
    header.is_linked_to_previous = False
    # Remove parágrafo vazio padrão, se houver texto vazio.
    for p in header.paragraphs:
        p.text = ""

    table = header.add_table(rows=1, cols=2, width=Cm(17.4))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(14.0)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)

    if logo_path:
        p_logo = table.cell(0, 0).paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        try:
            run.add_picture(str(logo_path), width=Cm(2.6))
        except Exception:
            set_font(run, size=10, bold=True)
            run.text = "TCE-RJ"
    else:
        p_logo = table.cell(0, 0).paragraphs[0]
        r = p_logo.add_run("TCE-RJ")
        set_font(r, size=16, bold=True, color="1F4E79")

    p_head = table.cell(0, 1).paragraphs[0]
    p_head.paragraph_format.space_after = Pt(0)
    lines = [
        ("TRIBUNAL DE CONTAS DO ESTADO DO RIO DE JANEIRO", True),
        ("SECRETARIA GERAL DE CONTROLE EXTERNO", False),
        ("SUBSECRETARIA DE CONTROLE DA RECEITA E POLÍTICAS DE CIDADANIA", False),
        ("COORDENADORIA SETORIAL DE AUDITORIA EM POLÍTICAS DE TECNOLOGIA DA INFORMAÇÃO", False),
    ]
    for i, (txt, bold) in enumerate(lines):
        if i:
            p_head.add_run().add_break()
        r = p_head.add_run(txt)
        set_font(r, size=9.5, bold=bold)

    p_line = header.add_paragraph()
    p_line.paragraph_format.space_before = Pt(2)
    p_line.paragraph_format.space_after = Pt(8)
    set_paragraph_bottom_border(p_line, size="8")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    set_font(r, size=8, color="808080")


def add_question_title(doc: Document, q: Question):
    text = clean_text("\n".join([ln for ln in q.text_lines if ln.strip()]))
    if not text:
        text = q.code
    # Quebra linhas longas preservando aparência de pergunta em negrito.
    add_text_paragraph(doc, text, bold=True, size=11, space_after=8, align=WD_ALIGN_PARAGRAPH.LEFT, style=resolve_style(doc, "Título 3", "Heading 3"))


def add_subgroup_title(doc: Document, subgroup: str):
    subgroup = clean_text(subgroup)
    if subgroup:
        add_text_paragraph(doc, subgroup, bold=True, size=14, space_after=8, style=resolve_style(doc, "Título 2", "Heading 2"))


def add_help(doc: Document, text: str):
    text = clean_text(text)
    if not text:
        return
    add_text_paragraph(doc, f"? {text}", size=10.5, color="C00000", space_after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY, style="Ajuda")


def add_adoption_explain(doc: Document, text: str):
    text = clean_text(text)
    if text:
        add_text_paragraph(
            doc,
            text,
            size=9,
            color="5B9BD5",
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            left_indent=4,
            name="Calibri",
            space_after=6,
        )


def add_adoption_evidence_text(doc: Document, text: str):
    text = clean_text(text)
    if text:
        add_adoption_evidence_callout(doc, text)


def add_adoption_evidence_callout(doc: Document, text: str):
    add_blank(doc, 2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(0.8)
    table.columns[1].width = Cm(17.4)

    icon_cell = table.cell(0, 0)
    text_cell = table.cell(0, 1)
    icon_cell.width = Cm(0.8)
    text_cell.width = Cm(17.4)
    icon_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for cell in (icon_cell, text_cell):
        set_cell_shading(cell, "EAF4FF")
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": "D6EAFB"},
            bottom={"val": "single", "sz": "4", "color": "D6EAFB"},
            left={"val": "single", "sz": "10", "color": "5B9BD5"},
            right={"val": "single", "sz": "4", "color": "D6EAFB"},
        )
        clear_cell(cell)

    p_icon = icon_cell.add_paragraph()
    p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_icon.paragraph_format.space_before = Pt(3)
    p_icon.paragraph_format.space_after = Pt(0)
    r_icon = p_icon.add_run("📄")
    set_font(r_icon, size=13, color="5B9BD5", name="Segoe UI Emoji")

    p_label = text_cell.add_paragraph()
    p_label.paragraph_format.space_before = Pt(3)
    p_label.paragraph_format.space_after = Pt(1)
    r_label = p_label.add_run("Evidência documental")
    set_font(r_label, bold=True, size=9, color="1F4E79", name="Calibri")

    p_body = text_cell.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after = Pt(6)
    r_body = p_body.add_run(text)
    set_font(r_body, size=9, color="35363F", name="Calibri")

    add_blank(doc, 4)


def add_prompt(doc: Document, text: str, indent=0.85):
    add_text_paragraph(doc, f"○ {clean_text(text)}", size=10, color="666666", left_indent=indent, space_after=2, style="ExplicaAlternativa")


def add_radio_option(doc: Document, text: str, indent=0.35):
    add_text_paragraph(doc, f"○  {clean_text(text)}", size=10.5, left_indent=indent, space_after=2, style="AlternativaTexto")


def add_checkbox_option(doc: Document, text: str, indent=0.35):
    add_text_paragraph(doc, f"☐  {clean_text(text)}", size=10.5, left_indent=indent, space_after=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY, style="AlternativaTexto")


def add_answer_line(doc: Document, label: str = "Resposta"):
    add_text_paragraph(doc, f"{label}: ________________________________________________________________", size=10, color="666666", left_indent=0.35, space_after=6, style="ExplicaAlternativa")


def clear_cell(cell):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


def add_cell_paragraph(cell, text: str = "", *, style: Optional[str] = None, align=None):
    clear_cell(cell)
    try:
        p = cell.add_paragraph(style=style)
    except KeyError:
        p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text:
        p.add_run(text)
    return p


def format_option_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(0.65)
    table.columns[1].width = Cm(17.6)
    for row in table.rows:
        row.cells[0].width = Cm(0.65)
        row.cells[1].width = Cm(17.6)
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "nil", "sz": "0", "color": "FFFFFF"},
                left={"val": "nil", "sz": "0", "color": "FFFFFF"},
                right={"val": "nil", "sz": "0", "color": "FFFFFF"},
            )


def add_option_row(table, text: str, *, symbol: str = "🔿", style: str = "AlternativaTexto"):
    row = table.add_row()
    format_option_table(table)
    p_symbol = add_cell_paragraph(row.cells[0], align=WD_ALIGN_PARAGRAPH.CENTER)
    p_symbol.add_run(symbol)
    p_text = add_cell_paragraph(row.cells[1], clean_text(text), style=style)
    return row, p_text


def add_explanation_row(table, text: str):
    row = table.add_row()
    format_option_table(table)
    add_cell_paragraph(row.cells[0])
    p_text = add_cell_paragraph(row.cells[1], style="ExplicaAlternativa")
    p_text.add_run("⮩ ")
    p_text.add_run(clean_text(text))
    return row, p_text


def add_option_table(doc: Document):
    table = doc.add_table(rows=0, cols=2)
    format_option_table(table)
    return table


def nested_question_text(q: Question) -> str:
    text = clean_text("\n".join([ln for ln in q.text_lines if ln.strip()])) or q.code
    parts = [text]
    if q.type in {"multi_text", "multitext"}:
        parts.extend(f"{clean_text(sq.text)}: ______________________________________________" for sq in q.subquestions)
    elif q.type in {"long", "textarea"}:
        parts.extend(["____________________________________________________________", "____________________________________________________________"])
    elif q.type in {"short", "text"}:
        parts.append("Resposta: ______________________________________________")
    elif q.type in {"multi", "multiple", "checkbox"}:
        parts.extend(f"☐  {clean_text(sq.text)}" for sq in (q.subquestions or q.options))
    elif q.type in {"single", "list", "radio"}:
        parts.extend(f"○  {clean_text(opt.text)}" for opt in q.options)
    elif q.type in {"upload", "file"}:
        parts.append("☐  Evidência documental a ser anexada no LimeSurvey.")
    return "\n".join(part for part in parts if part)


def options_for_question(survey: Survey, q: Question) -> List[Option]:
    if q.options:
        return q.options
    if q.scale and q.scale in survey.scales:
        return survey.scales[q.scale].options
    if q.type == "adoption":
        return ADOPTION_OPTIONS
    return []


def render_adoption(doc: Document, survey: Survey, q: Question, detail_dependencies: Optional[Dict[Tuple[str, str], List[Question]]] = None):
    start_question_page(doc)
    add_subgroup_title(doc, q.subgroup)
    add_question_title(doc, q)
    add_adoption_explain(doc, q.attrs.get("explain", ""))
    table = add_option_table(doc)
    for opt in ADOPTION_OPTIONS:
        add_option_row(table, opt.text)
        if opt.code == "naoap":
            for nsa in NSA_OPTIONS:
                add_option_row(table, '🔿 ' + nsa.text, symbol='')
                if nsa.code == "A":
                    add_explanation_row(table, "Indique que leis e/ou normas são essas:")
                elif nsa.code == "B":
                    add_explanation_row(table, "Identifique esses estudos:")
                else:
                    add_explanation_row(table, "Explique que razões são essas:")

    detail_value = q.attrs.get("detail", "true").lower()
    if detail_value not in {"n", "no", "false", "nao", "não", "0"} and q.detail_options:
        add_blank(doc, 4)
        add_text_paragraph(doc, "Visando explicitar melhor o grau de adoção do controle, marque abaixo uma ou mais opções que majoritariamente caracterizam sua organização:", style="AlternativaTexto")

        for opt in q.detail_options:
            add_text_paragraph(doc, f"☐  {opt.text}", style="AlternativaTexto")
            for depq in (detail_dependencies or {}).get((q.code, opt.code), []):
                texto_limpo = ''.join(depq.text_lines).strip()
                p = add_text_paragraph(doc, f"⮩ {texto_limpo}", style="ExplicaAlternativa")
                p.paragraph_format.left_indent = Inches(0.19)

            # Evidências de detalhe podem vir do XLSX ou de campos opcionais no futuro.
            if opt.evidence.lower() == "upload":
                add_explanation_row(table, opt.evidence_text or f"Anexe evidência documental referente ao item {opt.code}.")

    add_help(doc, q.help)
    add_adoption_evidence_text(doc, q.attrs.get("evidence_text", ""))
    add_blank(doc, 8)


def render_question(doc: Document, survey: Survey, q: Question, answer_dependencies: Optional[Dict[Tuple[str, str], List[Question]]] = None, detail_dependencies: Optional[Dict[Tuple[str, str], List[Question]]] = None):
    if q.type == "adoption":
        render_adoption(doc, survey, q, detail_dependencies=detail_dependencies)
        return

    start_question_page(doc)
    add_subgroup_title(doc, q.subgroup)
    add_question_title(doc, q)
    add_adoption_explain(doc, q.attrs.get("explain", ""))

    if q.type in {"single", "list", "radio"}:
        table = add_option_table(doc)
        for opt in options_for_question(survey, q):
            add_option_row(table, opt.text)
            for depq in (answer_dependencies or {}).get((q.code, opt.code), []):
                clean_depq = clean_text("\n".join([ln for ln in depq.text_lines if ln.strip()]))
                add_explanation_row(table, clean_depq)
    elif q.type in {"multi", "multiple", "checkbox"}:
        table = add_option_table(doc)
        if q.subquestions:
            for opt in q.subquestions:
                add_option_row(table, opt.text, symbol="☐")
        else:
            for opt in options_for_question(survey, q):
                add_option_row(table, opt.text, symbol="☐")
    elif q.type in {"multi_text", "multitext"}:
        for opt in q.subquestions:
            add_answer_line(doc, clean_text(opt.text))
    elif q.type in {"short", "text"}:
        add_answer_line(doc)
    elif q.type in {"long", "textarea"}:
        add_answer_line(doc)
        add_text_paragraph(doc, "________________________________________________________________", size=10, color="666666", left_indent=0.35, space_after=3, style="ExplicaAlternativa")
        add_text_paragraph(doc, "________________________________________________________________", size=10, color="666666", left_indent=0.35, space_after=6, style="ExplicaAlternativa")
    # elif q.type in {"upload", "file"}:
        # add_text_paragraph(doc, "☐  Evidência documental a ser anexada no LimeSurvey.", size=10, color="666666", left_indent=0.35, space_after=6, style="ExplicaAlternativa")
    elif q.type in {"array", "matrix"}:
        render_array_table(doc, survey, q)
    elif q.type in {"array_numbers", "array_number", "numeric_array", "array_numeros", "matriz_numerica", ":"}:
        render_array_table(doc, survey, q, cell_text=" ")
    else:
        add_answer_line(doc)

    add_help(doc, q.help)
    add_adoption_evidence_text(doc, q.attrs.get("evidence_text", ""))
    add_blank(doc, 8)


def render_array_table(doc: Document, survey: Survey, q: Question, cell_text: str = "○"):
    rows = q.subquestions or []
    cols = options_for_question(survey, q)
    if not rows or not cols:
        add_answer_line(doc)
        return
    table = doc.add_table(rows=len(rows) + 1, cols=len(cols) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = ""
    for i, col in enumerate(cols, start=1):
        hdr[i].text = clean_text(col.text)
        set_cell_shading(hdr[i], "D9EAF7")
    for r_idx, row in enumerate(rows, start=1):
        table.rows[r_idx].cells[0].text = clean_text(row.text)
        for c_idx in range(1, len(cols) + 1):
            table.rows[r_idx].cells[c_idx].text = cell_text
    for row_idx, row in enumerate(table.rows):
        is_header = row_idx == 0
        for cell_idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                if cell_idx == 0 and not is_header:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run, bold=is_header, size=10 if is_header else 9, name="Calibri")
    add_blank(doc, 8)




def simple_visible_dependency(expr: str) -> Optional[Tuple[str, str]]:
    """Retorna (codigo_da_questao_pai, valor) para expressões simples: q1011 == sim."""
    if not expr:
        return None
    expr = expr.strip()
    m = re.match(r"^([A-Za-z][A-Za-z0-9_]*)\s*==\s*['\"]?([A-Za-z0-9_\-]+)['\"]?$", expr)
    if m:
        return m.group(1), m.group(2)
    return None


def simple_detail_dependency(expr: str) -> Optional[Tuple[str, str]]:
    """Retorna (codigo_questao_base, subitem) para q1022ext.D == Y ou q1022ext[D] == Y."""
    if not expr:
        return None
    expr = expr.strip()
    m = re.match(r"^([A-Za-z][A-Za-z0-9_]*?)(?:\.|\[)([A-Za-z0-9_\-]+)\]?\s*==\s*['\"]?Y['\"]?$", expr, flags=re.IGNORECASE)
    if not m:
        return None
    parent = m.group(1)
    sub = m.group(2)
    base = parent[:-3] if parent.endswith("ext") else parent
    return base, sub


def build_dependency_maps(group: Group) -> Tuple[Dict[Tuple[str, str], List[Question]], Dict[Tuple[str, str], List[Question]], set[str]]:
    by_answer: Dict[Tuple[str, str], List[Question]] = {}
    by_detail: Dict[Tuple[str, str], List[Question]] = {}
    skipped: set[str] = set()
    codes = {q.code for q in group.questions}
    for q in group.questions:
        dep = simple_visible_dependency(q.visible_if)
        if dep and dep[0] in codes:
            by_answer.setdefault(dep, []).append(q)
            skipped.add(q.code)
            continue
        ddep = simple_detail_dependency(q.visible_if)
        if ddep and ddep[0] in codes:
            by_detail.setdefault(ddep, []).append(q)
            skipped.add(q.code)
            continue
    return by_answer, by_detail, skipped


def filter_survey_by_target(survey: Survey, target: str) -> Survey:
    return filter_target_survey(survey, target)


def render_nested_question(doc: Document, q: Question, indent: float = 0.85):
    text = clean_text("\n".join([ln for ln in q.text_lines if ln.strip()])) or q.code
    add_prompt(doc, text, indent=indent)
    if q.type in {"multi_text", "multitext"}:
        for sq in q.subquestions:
            add_text_paragraph(doc, f"{clean_text(sq.text)}: ______________________________________________", size=9.5, color="666666", left_indent=indent + 0.25, space_after=1, style="ExplicaAlternativa")
    elif q.type in {"long", "textarea"}:
        add_text_paragraph(doc, "____________________________________________________________", size=9.5, color="666666", left_indent=indent + 0.25, space_after=1, style="ExplicaAlternativa")
        add_text_paragraph(doc, "____________________________________________________________", size=9.5, color="666666", left_indent=indent + 0.25, space_after=2, style="ExplicaAlternativa")
    elif q.type in {"short", "text"}:
        add_text_paragraph(doc, "Resposta: ______________________________________________", size=9.5, color="666666", left_indent=indent + 0.25, space_after=2, style="ExplicaAlternativa")
    elif q.type in {"multi", "multiple", "checkbox"}:
        for sq in q.subquestions or q.options:
            add_text_paragraph(doc, f"☐  {clean_text(sq.text)}", size=9.5, color="666666", left_indent=indent + 0.25, space_after=1, style="AlternativaTexto")
    elif q.type in {"single", "list", "radio"}:
        for opt in q.options:
            add_text_paragraph(doc, f"○  {clean_text(opt.text)}", size=9.5, color="666666", left_indent=indent + 0.25, space_after=1, style="AlternativaTexto")
    elif q.type in {"upload", "file"}:
        add_text_paragraph(doc, "☐  Evidência documental a ser anexada no LimeSurvey.", size=9.5, color="666666", left_indent=indent + 0.25, space_after=2, style="ExplicaAlternativa")

def clear_document_body(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_docx(survey: Survey, output_path: Path, template_docx: Optional[Path] = None, logo_path_arg: Optional[Path] = None):
    if template_docx and template_docx.exists():
        doc = Document(template_docx)
        clear_document_body(doc)
    else:
        logo_path = logo_path_arg if logo_path_arg and logo_path_arg.exists() else extract_logo_from_template(template_docx)
        doc = Document()
        configure_document(doc, logo_path)

    # Título principal, se não for redundante com o primeiro grupo.
    if survey.title and "Questionário" in survey.title:
        add_text_paragraph(doc, survey.title.upper(), bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    for group_idx, group in enumerate(survey.groups):
        if group_idx > 0:
            doc.add_page_break()
        add_text_paragraph(doc, clean_text(group.title), bold=True, size=14, space_after=10, style=resolve_style(doc, "Título 1", "Heading 1"))
        for desc in group.description_lines:
            add_text_paragraph(doc, clean_text(desc), size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, style="TextoGrupo")
        if group.description_lines:
            add_blank(doc, 4)
        answer_deps, detail_deps, skipped = build_dependency_maps(group)
        for q in group.questions:
            if q.code in skipped:
                continue
            render_question(doc, survey, q, answer_dependencies=answer_deps, detail_dependencies=detail_deps)

    doc.save(output_path)


def main(argv: Optional[List[str]] = None):
    parser = argparse.ArgumentParser(description="Converte SurveyMD para DOCX de revisão/impressão.")
    parser.add_argument("input_md", type=Path, help="Arquivo .md SurveyMD de entrada")
    parser.add_argument("output_docx", type=Path, help="Arquivo .docx de saída")
    parser.add_argument("--template-docx", type=Path, default=None, help="DOCX de referência para extrair logotipo do cabeçalho")
    parser.add_argument("--logo", type=Path, default=None, help="Imagem de logotipo a ser usada no cabeçalho")
    args = parser.parse_args(argv)

    survey = parse_markdown(args.input_md)
    targets = survey_targets(survey)
    if targets:
        multi = len(targets) > 1
        for target in targets:
            target_survey = filter_survey_by_target(survey, target)
            output_path = target_output_path(args.output_docx, target) if multi else args.output_docx
            build_docx(target_survey, output_path, template_docx=args.template_docx, logo_path_arg=args.logo)
            print(f"DOCX gerado: {output_path}")
        return
    build_docx(survey, args.output_docx, template_docx=args.template_docx, logo_path_arg=args.logo)
    print(f"DOCX gerado: {args.output_docx}")


if __name__ == "__main__":
    main()
