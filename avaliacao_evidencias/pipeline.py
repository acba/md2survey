from __future__ import annotations

import argparse
import difflib
import hashlib
import io
import json
import datetime as dt
import os
import re
import shutil
import sys
import tempfile
import time
import unicodedata
import zipfile
from urllib.parse import unquote
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Iterable

import md2lss
from openpyxl import Workbook, load_workbook

from .providers_ai_service import executar_provider


@dataclass(frozen=True)
class AnaliseCandidata:
    auditado: str
    coluna_evidencia: str
    nome_original_evidencia: str
    upload: dict[str, Any]
    resposta_id: Any = None
    evidence_index: int | None = None
    erro: str = ""


@dataclass(frozen=True)
class ResolucaoEvidencia:
    caminho: Path | None
    nome_decodificado: str
    erro: str = ""


@dataclass(frozen=True)
class ItemAfirmado:
    codigo: str
    texto: str
    afirmacao: str


@dataclass(frozen=True)
class QuestaoContexto:
    codigo: str
    tipo: str
    texto: str
    itens: dict[str, str]


@dataclass(frozen=True)
class ContextoQuestionario:
    questoes: dict[str, QuestaoContexto]


@dataclass(frozen=True)
class ChecklistResolvido:
    nome: str
    caminho: Path | None
    conteudo: str
    hash_conteudo: str
    erro: str = ""


@dataclass(frozen=True)
class PromptResolvido:
    nome: str
    caminho: Path | None
    conteudo: str
    hash_conteudo: str
    erro: str = ""


@dataclass(frozen=True)
class PacoteEvidencia:
    caminho: Path
    tipo: str
    documentos: list[dict[str, Any]]
    inventario: list[str]
    erro: str = ""


REMOTE_PROVIDERS = {"gemini", "openrouter"}


def log_event(event: str, message: str, *, quiet: bool = False, level: str = "info", **fields: Any) -> None:
    if quiet:
        return
    payload = {
        "ts": dt.datetime.now(dt.timezone.utc).isoformat(),
        "level": level,
        "event": event,
        "message": message,
        **fields,
    }
    print(json.dumps(payload, ensure_ascii=False, default=str), file=sys.stdout, flush=True)


def validar_rpm(valor: str) -> int:
    try:
        rpm = int(valor)
    except ValueError as exc:
        raise argparse.ArgumentTypeError("--rpm deve ser um numero inteiro maior ou igual a zero") from exc
    if rpm < 0:
        raise argparse.ArgumentTypeError("--rpm deve ser maior ou igual a zero")
    return rpm


@dataclass
class RequestsPerMinuteLimiter:
    rpm: int
    clock: Callable[[], float] = time.monotonic
    sleeper: Callable[[float], None] = time.sleep
    last_started_at: float | None = None

    def wait_seconds(self) -> float:
        if self.rpm <= 0 or self.last_started_at is None:
            return 0.0
        elapsed = self.clock() - self.last_started_at
        return max(0.0, (60.0 / self.rpm) - elapsed)

    def wait_and_mark(self, wait_seconds: float | None = None) -> float:
        if wait_seconds is None:
            wait_seconds = self.wait_seconds()
        if wait_seconds > 0:
            self.sleeper(wait_seconds)
        self.last_started_at = self.clock()
        return wait_seconds


def coluna_evidencia(nome: str) -> bool:
    return "evi" in nome and not nome.endswith("[filecount]")


def _parse_upload(valor: Any) -> tuple[dict[str, Any] | None, str]:
    if valor in (None, ""):
        return None, ""
    if isinstance(valor, str):
        try:
            parsed = json.loads(valor)
        except json.JSONDecodeError as exc:
            return None, f"metadados de upload invalidos: {exc.msg}"
    else:
        parsed = valor
    if not isinstance(parsed, list):
        return None, "metadados de upload devem ser uma lista"
    if len(parsed) != 1:
        return None, "coluna de evidencia deve conter exatamente um arquivo"
    upload = parsed[0]
    if not isinstance(upload, dict):
        return None, "metadado de upload deve ser um objeto"
    name = upload.get("name")
    if not isinstance(name, str) or not name:
        return None, "metadado de upload sem atributo name"
    return upload, ""


def _rows_from_xlsx(caminho_xlsx: Path) -> Iterable[dict[str, Any]]:
    workbook = load_workbook(caminho_xlsx, read_only=True, data_only=True)
    try:
        sheet = workbook.active
        rows = sheet.iter_rows(values_only=True)
        headers = [str(value) if value is not None else "" for value in next(rows)]
        for values in rows:
            yield dict(zip(headers, values))
    finally:
        workbook.close()


def inventariar_analises(
    caminho_xlsx: str | Path,
    raiz_evidencias: str | Path,
    caminho_questionario: str | Path,
    *,
    include_unsubmitted: bool = False,
) -> list[AnaliseCandidata]:
    del raiz_evidencias, caminho_questionario
    linhas = list(_rows_from_xlsx(Path(caminho_xlsx)))
    if not linhas:
        return []
    colunas = list(linhas[0].keys())
    colunas_evidencia = [col for col in colunas if coluna_evidencia(col)]
    analises: list[AnaliseCandidata] = []
    for linha in linhas:
        if not include_unsubmitted and not linha.get("submitdate"):
            continue
        auditado = str(linha.get("firstname") or "").strip()
        for evidence_index, coluna in enumerate(colunas_evidencia, start=1):
            upload, erro = _parse_upload(linha.get(coluna))
            if upload is None and not erro:
                continue
            analises.append(
                AnaliseCandidata(
                    auditado=auditado,
                    coluna_evidencia=coluna,
                    nome_original_evidencia=upload.get("name", "") if upload else "",
                    upload=upload or {},
                    resposta_id=linha.get("id"),
                    evidence_index=evidence_index,
                    erro=erro,
                )
            )
    return analises


def resolver_evidencia(
    auditado: str,
    raiz_evidencias: str | Path,
    upload: dict[str, Any],
    *,
    resposta_id: Any = None,
    evidence_index: int | None = None,
) -> ResolucaoEvidencia:
    nome_original = upload.get("name")
    if not isinstance(nome_original, str) or not nome_original:
        return ResolucaoEvidencia(caminho=None, nome_decodificado="", erro="metadado de upload sem atributo name")
    nome_decodificado = unquote(nome_original)
    auditado_dir = Path(raiz_evidencias) / auditado
    caminho = auditado_dir / nome_decodificado
    if not caminho.is_file():
        exportado = _resolver_evidencia_exportada_limesurvey(
            auditado_dir,
            nome_decodificado,
            resposta_id=resposta_id,
            evidence_index=evidence_index,
        )
        if exportado:
            return ResolucaoEvidencia(caminho=exportado, nome_decodificado=nome_decodificado)
        return ResolucaoEvidencia(
            caminho=None,
            nome_decodificado=nome_decodificado,
            erro=f"arquivo de evidencia nao encontrado: {caminho}",
        )
    return ResolucaoEvidencia(caminho=caminho, nome_decodificado=nome_decodificado)


def _resolver_evidencia_exportada_limesurvey(
    auditado_dir: Path,
    nome_decodificado: str,
    *,
    resposta_id: Any = None,
    evidence_index: int | None = None,
) -> Path | None:
    if not auditado_dir.is_dir():
        return None
    candidatos = [path for path in auditado_dir.iterdir() if path.is_file()]
    if not candidatos:
        return None

    if resposta_id is not None and evidence_index is not None:
        prefixo = _prefixo_exportacao_limesurvey(resposta_id, evidence_index)
        if prefixo:
            candidatos_prefixo = [path for path in candidatos if path.name.startswith(prefixo)]
            match = _melhor_match_nome_exportado(nome_decodificado, candidatos_prefixo)
            if match:
                return match

    return _melhor_match_nome_exportado(nome_decodificado, candidatos)


def _prefixo_exportacao_limesurvey(resposta_id: Any, evidence_index: int) -> str:
    try:
        return f"{int(resposta_id):05d}_{int(evidence_index):02d}_"
    except (TypeError, ValueError):
        return ""


def _melhor_match_nome_exportado(nome_decodificado: str, candidatos: list[Path]) -> Path | None:
    nome_key, suffix = _nome_exportado_key(nome_decodificado)
    matches: list[tuple[float, str, Path]] = []
    for candidato in candidatos:
        candidato_sem_prefixo = _remover_prefixo_exportacao_limesurvey(candidato.name)
        candidato_key, candidato_suffix = _nome_exportado_key(candidato_sem_prefixo)
        if candidato_suffix != suffix:
            continue
        if candidato_key == nome_key:
            score = 1.0
        else:
            score = difflib.SequenceMatcher(None, nome_key, candidato_key).ratio()
        if score >= 0.92:
            matches.append((score, candidato.name, candidato))
    if not matches:
        return None
    matches.sort(key=lambda item: (-item[0], item[1]))
    return matches[0][2]


def _remover_prefixo_exportacao_limesurvey(nome: str) -> str:
    return re.sub(r"^\d+_\d+_", "", nome)


def _nome_exportado_key(nome: str) -> tuple[str, str]:
    path = Path(nome)
    stem = unicodedata.normalize("NFKD", path.stem.casefold())
    stem = "".join(char for char in stem if not unicodedata.combining(char))
    stem = re.sub(r"[^a-z0-9]+", "-", stem)
    stem = re.sub(r"-+", "-", stem).strip("-")
    return stem, path.suffix.casefold()


def carregar_contexto_questionario(caminho_questionario: str | Path) -> ContextoQuestionario:
    survey = md2lss.parse_markdown(Path(caminho_questionario))
    questoes: dict[str, QuestaoContexto] = {}
    for group in survey.groups:
        for question in group.questions:
            if question.type == "upload":
                continue
            itens = {option.code: option.text for option in (question.subquestions or question.alternatives)}
            questoes[question.code] = QuestaoContexto(
                codigo=question.code,
                tipo=question.type,
                texto=question.text(),
                itens=itens,
            )
    return ContextoQuestionario(questoes=questoes)


def _base_coluna_evidencia(coluna_evidencia: str) -> tuple[str, str | None]:
    marker = coluna_evidencia.find("evi")
    base = coluna_evidencia[:marker]
    suffix = coluna_evidencia[marker + 3 :]
    return base, suffix or None


def _valor_afirmativo(valor: Any) -> bool:
    if isinstance(valor, str):
        return valor.strip().casefold() in {"sim", "y", "yes", "true", "1"}
    return valor is True or valor == 1


ADOPTION_VALUES = {"naoad", "adfor", "admen", "adpar", "admai", "naoap"}


def selecionar_itens_afirmados(
    contexto: ContextoQuestionario,
    coluna_evidencia: str,
    resposta: dict[str, Any],
) -> list[ItemAfirmado]:
    base, item_especifico = _base_coluna_evidencia(coluna_evidencia)
    questao = contexto.questoes.get(base)
    if not questao:
        return []

    if item_especifico:
        valor = resposta.get(f"{base}[{item_especifico}]")
        if _valor_afirmativo(valor):
            return [
                ItemAfirmado(
                    codigo=f"{base}[{item_especifico}]",
                    texto=questao.itens.get(item_especifico, item_especifico),
                    afirmacao=str(valor),
                )
            ]
        return []

    valor_base = resposta.get(base)
    if questao.tipo in {"single", "adoption"} and valor_base in {"adpar", "admai"}:
        itens = [ItemAfirmado(codigo=base, texto=questao.texto, afirmacao=str(valor_base))]
        prefixo_ext = f"{base}ext["
        detalhe = contexto.questoes.get(f"{base}ext")
        for chave, valor in resposta.items():
            if chave.startswith(prefixo_ext) and chave.endswith("]") and valor == "Y":
                codigo_item = chave[len(prefixo_ext) : -1]
                itens.append(
                    ItemAfirmado(
                        codigo=chave,
                        texto=(detalhe.itens if detalhe else {}).get(codigo_item, codigo_item),
                        afirmacao="Y",
                    )
                )
        return itens

    if questao.tipo == "single" and valor_base in ADOPTION_VALUES:
        return []

    if questao.tipo == "single" and valor_base not in (None, ""):
        valor_codigo = str(valor_base)
        return [
            ItemAfirmado(
                codigo=f"{base}[{valor_codigo}]",
                texto=questao.itens.get(valor_codigo, questao.texto),
                afirmacao=valor_codigo,
            )
        ]

    itens = []
    for codigo_item, texto in questao.itens.items():
        chave = f"{base}[{codigo_item}]"
        valor = resposta.get(chave)
        if _valor_afirmativo(valor):
            itens.append(ItemAfirmado(codigo=chave, texto=texto, afirmacao=str(valor)))
    return itens


def resolver_checklist(checklists_dir: str | Path, coluna_evidencia: str) -> ChecklistResolvido:
    base, item_especifico = _base_coluna_evidencia(coluna_evidencia)
    raiz = Path(checklists_dir)
    candidatos = []
    if item_especifico:
        candidatos.append(raiz / f"{base}_{item_especifico}.md")
    candidatos.append(raiz / f"{base}.md")
    for caminho in candidatos:
        if caminho.is_file():
            conteudo = caminho.read_text(encoding="utf-8")
            digest = hashlib.sha256(conteudo.encode("utf-8")).hexdigest()
            return ChecklistResolvido(
                nome=caminho.name,
                caminho=caminho,
                conteudo=conteudo,
                hash_conteudo=digest,
            )
    return ChecklistResolvido(
        nome="",
        caminho=None,
        conteudo="",
        hash_conteudo="",
        erro=f"checklist de analise nao encontrado para {coluna_evidencia}",
    )


def resolver_prompt(prompts_dir: str | Path, coluna_evidencia: str) -> PromptResolvido:
    base, item_especifico = _base_coluna_evidencia(coluna_evidencia)
    raiz = Path(prompts_dir)
    candidatos = []
    if item_especifico:
        candidatos.append(raiz / f"{base}_{item_especifico}.md")
    candidatos.append(raiz / f"{base}.md")
    for caminho in candidatos:
        if caminho.is_file():
            conteudo = caminho.read_text(encoding="utf-8")
            digest = hashlib.sha256(conteudo.encode("utf-8")).hexdigest()
            return PromptResolvido(
                nome=caminho.name,
                caminho=caminho,
                conteudo=conteudo,
                hash_conteudo=digest,
            )
    return PromptResolvido(
        nome="",
        caminho=None,
        conteudo="",
        hash_conteudo="",
        erro=f"prompt de analise nao encontrado para {coluna_evidencia}",
    )


def hash_arquivo(caminho: str | Path) -> str:
    digest = hashlib.sha256()
    with Path(caminho).open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _zip_member_safe(name: str) -> bool:
    path = Path(name)
    return not path.is_absolute() and ".." not in path.parts


def normalizar_evidencia(caminho: str | Path) -> PacoteEvidencia:
    path = Path(caminho)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        try:
            texto = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            texto = path.read_text(encoding="latin-1")
        return PacoteEvidencia(
            caminho=path,
            tipo=suffix.lstrip("."),
            documentos=[{"nome": path.name, "texto": texto}],
            inventario=[path.name],
        )
    if suffix == ".zip":
        documentos: list[dict[str, Any]] = []
        try:
            with zipfile.ZipFile(path) as archive:
                names = archive.namelist()
                for name in names:
                    if not _zip_member_safe(name):
                        return PacoteEvidencia(
                            caminho=path,
                            tipo="zip",
                            documentos=[],
                            inventario=names,
                            erro=f"zip contem path traversal: {name}",
                        )
                for name in names:
                    if name.endswith("/"):
                        continue
                    suffix_interno = Path(name).suffix.lower()
                    if suffix_interno in {".txt", ".md", ".csv"}:
                        data = archive.read(name)
                        try:
                            texto = data.decode("utf-8")
                        except UnicodeDecodeError:
                            texto = data.decode("latin-1")
                        documentos.append({"nome": name, "texto": texto})
                    elif suffix_interno == ".pdf":
                        pdf_docs, erro_pdf = _extrair_texto_pdf_bytes(name, archive.read(name))
                        if pdf_docs:
                            documentos.extend(pdf_docs)
                        else:
                            documentos.append({"nome": name, "erro": erro_pdf})
                    else:
                        documentos.append({"nome": name, "nao_suportado": True})
                return PacoteEvidencia(caminho=path, tipo="zip", documentos=documentos, inventario=names)
        except zipfile.BadZipFile:
            return PacoteEvidencia(caminho=path, tipo="zip", documentos=[], inventario=[], erro="zip invalido")
    if suffix == ".pdf":
        documentos, erro = _extrair_texto_pdf(path)
        return PacoteEvidencia(
            caminho=path,
            tipo="pdf",
            documentos=documentos,
            inventario=[path.name],
            erro=erro,
        )
    if suffix == ".xlsx":
        workbook = load_workbook(path, read_only=True, data_only=True)
        try:
            documentos = []
            inventario = []
            for sheet in workbook.worksheets:
                inventario.append(sheet.title)
                linhas = []
                for row in sheet.iter_rows(values_only=True):
                    valores = ["" if value is None else str(value) for value in row]
                    if any(valores):
                        linhas.append("\t".join(valores))
                documentos.append({"nome": sheet.title, "texto": "\n".join(linhas)})
            return PacoteEvidencia(
                caminho=path,
                tipo="xlsx",
                documentos=documentos,
                inventario=inventario,
            )
        finally:
            workbook.close()
    if suffix == ".docx":
        texto, erro = _extrair_texto_docx(path)
        if erro:
            return PacoteEvidencia(
                caminho=path,
                tipo="docx",
                documentos=[],
                inventario=[path.name],
                erro=erro,
            )
        return PacoteEvidencia(
            caminho=path,
            tipo="docx",
            documentos=[{"nome": path.name, "texto": texto}],
            inventario=[path.name],
        )
    return PacoteEvidencia(
        caminho=path,
        tipo=suffix.lstrip(".") or "desconhecido",
        documentos=[],
        inventario=[path.name],
        erro=f"tipo de evidencia nao suportado: {suffix or path.name}",
    )


EXTENSOES_UPLOAD_DIRETO = {".pdf", ".txt", ".md", ".csv", ".xlsx", ".png", ".jpg", ".jpeg"}
EXTENSOES_UPLOAD_PREPARAVEIS = EXTENSOES_UPLOAD_DIRETO | {".docx"}


def _extrair_texto_docx(caminho: str | Path) -> tuple[str, str]:
    try:
        from docx import Document
    except ModuleNotFoundError:
        return "", "python-docx nao instalado para normalizar DOCX"
    document = Document(caminho)
    partes = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            partes.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(partes), ""


def _extrair_texto_pdf_reader(nome: str, reader: Any) -> tuple[list[dict[str, Any]], str]:
    documentos: list[dict[str, Any]] = []
    for index, page in enumerate(reader.pages, start=1):
        texto = (page.extract_text() or "").strip()
        if texto:
            documentos.append({"nome": nome, "pagina": index, "texto": texto})
    if not documentos:
        return [], "pdf sem texto extraivel"
    return documentos, ""


def _extrair_texto_pdf(caminho: str | Path) -> tuple[list[dict[str, Any]], str]:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        return [], "pypdf nao instalado para normalizar PDF"
    try:
        path = Path(caminho)
        return _extrair_texto_pdf_reader(path.name, PdfReader(path))
    except Exception as exc:
        return [], f"erro ao normalizar PDF: {exc}"


def _extrair_texto_pdf_bytes(nome: str, data: bytes) -> tuple[list[dict[str, Any]], str]:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        return [], "pypdf nao instalado para normalizar PDF"
    try:
        return _extrair_texto_pdf_reader(nome, PdfReader(io.BytesIO(data)))
    except Exception as exc:
        return [], f"erro ao normalizar PDF: {exc}"


def _slug_ascii(valor: str, *, fallback: str = "evidencia", max_len: int = 80) -> str:
    normalizado = unicodedata.normalize("NFKD", valor)
    ascii_text = normalizado.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^A-Za-z0-9._-]+", "-", ascii_text).strip("._-").lower()
    slug = re.sub(r"-{2,}", "-", slug)
    if not slug:
        slug = fallback
    return slug[:max_len].strip("._-") or fallback


def _nome_upload_seguro(nome_original: str, sufixo: str) -> str:
    suffix = sufixo.lower()
    stem = Path(nome_original).stem or "evidencia"
    slug = _slug_ascii(stem)
    digest = hashlib.sha256(nome_original.encode("utf-8", errors="ignore")).hexdigest()[:10]
    return f"{slug}-{digest}{suffix}"


def _caminho_unico(destino: Path, nome: str) -> Path:
    candidato = destino / nome
    if not candidato.exists():
        return candidato
    stem = candidato.stem
    suffix = candidato.suffix
    contador = 2
    while True:
        proximo = destino / f"{stem}-{contador}{suffix}"
        if not proximo.exists():
            return proximo
        contador += 1


def _copiar_upload_seguro(origem: Path, destino: Path, nome_original: str | None = None) -> Path:
    destino.mkdir(parents=True, exist_ok=True)
    nome = _nome_upload_seguro(nome_original or origem.name, origem.suffix)
    target = _caminho_unico(destino, nome)
    shutil.copyfile(origem, target)
    return target


def _gravar_upload_texto(destino: Path, nome_original: str, texto: str) -> Path:
    destino.mkdir(parents=True, exist_ok=True)
    nome = _nome_upload_seguro(nome_original, ".txt")
    target = _caminho_unico(destino, nome)
    target.write_text(texto, encoding="utf-8")
    return target


def _preparar_docx_para_upload(caminho: Path, destino: Path, nome_original: str | None = None) -> Path | None:
    texto, erro = _extrair_texto_docx(caminho)
    if erro:
        return None
    return _gravar_upload_texto(destino, nome_original or caminho.name, texto)


def arquivos_compativeis_upload(caminho: str | Path, destino_zip: str | Path | None = None) -> list[str]:
    path = Path(caminho)
    if path.is_file() and path.suffix.lower() in EXTENSOES_UPLOAD_PREPARAVEIS and path.suffix.lower() != ".zip":
        if destino_zip is None:
            return [str(path)] if path.suffix.lower() in EXTENSOES_UPLOAD_DIRETO else []
        destino = Path(destino_zip)
        if path.suffix.lower() == ".docx":
            preparado = _preparar_docx_para_upload(path, destino)
            return [str(preparado)] if preparado else []
        return [str(_copiar_upload_seguro(path, destino))]
    if path.suffix.lower() == ".zip" and destino_zip is not None:
        destino = Path(destino_zip)
        destino.mkdir(parents=True, exist_ok=True)
        arquivos = []
        with zipfile.ZipFile(path) as archive:
            for name in archive.namelist():
                if name.endswith("/") or not _zip_member_safe(name):
                    continue
                member = Path(name)
                suffix = member.suffix.lower()
                if suffix not in EXTENSOES_UPLOAD_PREPARAVEIS:
                    continue
                if suffix == ".docx":
                    staging_dir = destino / "_docx_sources"
                    staging_dir.mkdir(parents=True, exist_ok=True)
                    staging = _caminho_unico(staging_dir, _nome_upload_seguro(name, ".docx"))
                    staging.write_bytes(archive.read(name))
                    preparado = _preparar_docx_para_upload(staging, destino, name)
                    if preparado:
                        arquivos.append(str(preparado))
                    continue
                target = _caminho_unico(destino, _nome_upload_seguro(name, suffix))
                target.write_bytes(archive.read(name))
                arquivos.append(str(target))
        return arquivos
    return []


def calcular_identidade_analise(
    *,
    auditado: str,
    coluna_evidencia: str,
    nome_original_evidencia: str,
    hash_conteudo: str,
    provider: str,
    model: str,
    prompt_hash: str = "",
    checklist_hash: str = "",
    prompt_version: str,
) -> str:
    artifact_hash = prompt_hash or checklist_hash
    payload = {
        "auditado": auditado,
        "coluna_evidencia": coluna_evidencia,
        "nome_original_evidencia": nome_original_evidencia,
        "hash_conteudo": hash_conteudo,
        "provider": provider,
        "model": model,
        "prompt_hash": artifact_hash,
        "prompt_version": prompt_version,
    }
    encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def gravar_registro_analise(checkpoint: str | Path, registro: dict[str, Any]) -> None:
    path = Path(checkpoint)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as file:
        file.write(json.dumps(registro, ensure_ascii=False, sort_keys=True, default=str))
        file.write("\n")


def carregar_registros_analise(checkpoint: str | Path) -> dict[str, dict[str, Any]]:
    path = Path(checkpoint)
    if not path.is_file():
        return {}
    registros: dict[str, dict[str, Any]] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        registro = json.loads(line)
        identity = registro.get("identity")
        if isinstance(identity, str):
            registros[identity] = registro
    return registros


def deve_processar_identidade(
    registros: dict[str, dict[str, Any]],
    identity: str,
    *,
    skip_errors: bool = False,
) -> bool:
    registro = registros.get(identity)
    if not registro:
        return True
    status = registro.get("status")
    if status == "completed":
        return False
    if status == "error" and skip_errors:
        return False
    return True


def _join_value(value: Any) -> str:
    if isinstance(value, list):
        return "; ".join(str(item) for item in value)
    if value is None:
        return ""
    return str(value)


def gerar_relatorio_conformidade(checkpoint: str | Path, destino: str | Path) -> int:
    registros = carregar_registros_analise(checkpoint)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Conformidade"
    headers = [
        "auditado",
        "questao",
        "item",
        "afirmacao_auditado",
        "estado",
        "justificativa",
        "lacunas",
        "referencias",
        "evidencia",
        "provider",
        "model",
        "data_analise",
        "status_revisao_humana",
    ]
    sheet.append(headers)
    total_linhas = 0
    for registro in registros.values():
        result = registro.get("result") if isinstance(registro.get("result"), dict) else {}
        conclusoes = result.get("conclusoes") if isinstance(result, dict) else None
        if not conclusoes and registro.get("status") == "error":
            conclusoes = [
                {
                    "item_codigo": "",
                    "afirmacao_auditado": "",
                    "estado": "erro",
                    "justificativa": registro.get("error", "Erro de analise"),
                    "lacunas": [],
                    "arquivos_referenciados": [],
                }
            ]
        for conclusao in conclusoes or []:
            total_linhas += 1
            sheet.append(
                [
                    registro.get("auditado", ""),
                    registro.get("questao", ""),
                    conclusao.get("item_codigo", ""),
                    conclusao.get("afirmacao_auditado", ""),
                    conclusao.get("estado", ""),
                    conclusao.get("justificativa", ""),
                    _join_value(conclusao.get("lacunas")),
                    _join_value(conclusao.get("arquivos_referenciados")),
                    registro.get("evidencia", ""),
                    registro.get("provider", ""),
                    registro.get("model", ""),
                    registro.get("finished_at", ""),
                    None,
                ]
            )
    destino_path = Path(destino)
    destino_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(destino_path)
    return total_linhas


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Pre-analisa evidencias enviadas por auditados.")
    parser.add_argument("respostas")
    parser.add_argument("evidencias")
    parser.add_argument("--questionario", required=True)
    parser.add_argument("--provider", default="fake")
    parser.add_argument("--model", default="fake")
    parser.add_argument("--prompts-dir", default=None, help="Diretorio com Prompts de analise por questao.")
    parser.add_argument("--checklists-dir", default=None, help="Alias legado para --prompts-dir.")
    parser.add_argument("--out-dir", default=".saida_analise")
    parser.add_argument("--prompt-version", default="v1")
    parser.add_argument(
        "--rpm",
        type=validar_rpm,
        default=12,
        help="Limita requests por minuto para providers remotos; padrao 12, 0 desativa.",
    )
    parser.add_argument("--skip-errors", action="store_true")
    parser.add_argument(
        "--include-unsubmitted",
        action="store_true",
        help="Inclui respostas sem submitdate. Por padrao, somente respostas submetidas sao analisadas.",
    )
    parser.add_argument("--list-only", action="store_true")
    parser.add_argument("--quiet", action="store_true", help="Nao emite logs estruturados durante a execucao.")
    args = parser.parse_args(argv)
    prompts_dir = args.prompts_dir or args.checklists_dir or "checklists"
    rate_limiter = RequestsPerMinuteLimiter(args.rpm)
    log_event(
        "pipeline_started",
        "Inicio da avaliacao de evidencias.",
        quiet=args.quiet,
        respostas=args.respostas,
        evidencias=args.evidencias,
        questionario=args.questionario,
        prompts_dir=prompts_dir,
        provider=args.provider,
        model=args.model,
        out_dir=args.out_dir,
        prompt_version=args.prompt_version,
        rpm=args.rpm,
        include_unsubmitted=args.include_unsubmitted,
        skip_errors=args.skip_errors,
        list_only=args.list_only,
    )
    analises = inventariar_analises(
        args.respostas,
        args.evidencias,
        args.questionario,
        include_unsubmitted=args.include_unsubmitted,
    )
    log_event(
        "inventory_completed",
        "Inventario de evidencias concluido.",
        quiet=args.quiet,
        total_analises=len(analises),
        analises_com_erro=sum(1 for analise in analises if analise.erro),
        auditados=sorted({analise.auditado for analise in analises if analise.auditado}),
    )
    if args.list_only:
        log_event(
            "list_only_started",
            "Listando analises candidatas sem processar evidencias.",
            quiet=args.quiet,
            total_analises=len(analises),
        )
        for analise in analises:
            print(json.dumps(analise.__dict__, ensure_ascii=False, default=str))
        log_event(
            "pipeline_finished",
            "Execucao finalizada em modo list-only.",
            quiet=args.quiet,
            total_analises=len(analises),
        )
        return 0
    contexto = carregar_contexto_questionario(args.questionario)
    out_dir = Path(args.out_dir)
    checkpoint = out_dir / "analyses.jsonl"
    registros = carregar_registros_analise(checkpoint)
    log_event(
        "checkpoint_loaded",
        "Checkpoint carregado para deduplicacao.",
        quiet=args.quiet,
        checkpoint=str(checkpoint),
        registros=len(registros),
    )
    total_processadas = 0
    total_puladas = 0
    total_erros = 0
    total_concluidas = 0
    for index, analise in enumerate(analises, start=1):
        questao_base, _ = _base_coluna_evidencia(analise.coluna_evidencia)
        base_log = {
            "index": index,
            "total": len(analises),
            "auditado": analise.auditado,
            "questao": questao_base,
            "coluna_evidencia": analise.coluna_evidencia,
            "evidencia": analise.nome_original_evidencia,
        }
        log_event(
            "analysis_started",
            "Analise candidata iniciada.",
            quiet=args.quiet,
            **base_log,
        )
        if analise.erro:
            identity = hashlib.sha256(
                json.dumps(analise.__dict__, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
            ).hexdigest()
            if not deve_processar_identidade(registros, identity, skip_errors=args.skip_errors):
                total_puladas += 1
                log_event(
                    "analysis_skipped",
                    "Analise com erro de inventario ignorada por checkpoint.",
                    quiet=args.quiet,
                    identity=identity,
                    reason="checkpoint",
                    **base_log,
                )
            else:
                gravar_registro_analise(
                    checkpoint,
                    {
                        "identity": identity,
                        "status": "error",
                        "auditado": analise.auditado,
                        "questao": questao_base,
                        "coluna_evidencia": analise.coluna_evidencia,
                        "evidencia": analise.nome_original_evidencia,
                        "provider": args.provider,
                        "model": args.model,
                        "error": analise.erro,
                        "finished_at": dt.datetime.now(dt.timezone.utc).isoformat(),
                    },
                )
                registros[identity] = {"identity": identity, "status": "error"}
                total_processadas += 1
                total_erros += 1
                log_event(
                    "analysis_recorded_error",
                    "Erro de inventario registrado no checkpoint.",
                    quiet=args.quiet,
                    level="error",
                    identity=identity,
                    error=analise.erro,
                    **base_log,
                )
            continue
        resolucao = resolver_evidencia(
            analise.auditado,
            args.evidencias,
            analise.upload,
            resposta_id=analise.resposta_id,
            evidence_index=analise.evidence_index,
        )
        log_event(
            "evidence_resolved",
            "Resolucao do arquivo de evidencia concluida.",
            quiet=args.quiet,
            level="error" if resolucao.erro else "info",
            caminho=str(resolucao.caminho) if resolucao.caminho else "",
            nome_decodificado=resolucao.nome_decodificado,
            error=resolucao.erro,
            **base_log,
        )
        prompt = resolver_prompt(prompts_dir, analise.coluna_evidencia)
        log_event(
            "prompt_resolved",
            "Resolucao do prompt de analise concluida.",
            quiet=args.quiet,
            level="error" if prompt.erro else "info",
            prompt=prompt.nome,
            prompt_hash=prompt.hash_conteudo,
            error=prompt.erro,
            **base_log,
        )
        if resolucao.erro:
            hash_conteudo = ""
        else:
            hash_conteudo = hash_arquivo(resolucao.caminho)
        identity = calcular_identidade_analise(
            auditado=analise.auditado,
            coluna_evidencia=analise.coluna_evidencia,
            nome_original_evidencia=analise.nome_original_evidencia,
            hash_conteudo=hash_conteudo,
            provider=args.provider,
            model=args.model,
            prompt_hash=prompt.hash_conteudo,
            prompt_version=args.prompt_version,
        )
        if not deve_processar_identidade(registros, identity, skip_errors=args.skip_errors):
            total_puladas += 1
            log_event(
                "analysis_skipped",
                "Analise ignorada por ja existir no checkpoint.",
                quiet=args.quiet,
                identity=identity,
                reason="checkpoint",
                **base_log,
            )
            continue
        erro = resolucao.erro or prompt.erro
        if erro:
            gravar_registro_analise(
                checkpoint,
                {
                    "identity": identity,
                    "status": "error",
                    "auditado": analise.auditado,
                    "questao": questao_base,
                    "coluna_evidencia": analise.coluna_evidencia,
                    "evidencia": analise.nome_original_evidencia,
                    "provider": args.provider,
                    "model": args.model,
                    "error": erro,
                    "finished_at": dt.datetime.now(dt.timezone.utc).isoformat(),
                },
            )
            registros[identity] = {"identity": identity, "status": "error"}
            total_processadas += 1
            total_erros += 1
            log_event(
                "analysis_recorded_error",
                "Erro antes da chamada ao provider registrado no checkpoint.",
                quiet=args.quiet,
                level="error",
                identity=identity,
                error=erro,
                **base_log,
            )
            continue
        itens = selecionar_itens_afirmados(contexto, analise.coluna_evidencia, _linha_por_id(args.respostas, analise.resposta_id))
        log_event(
            "items_selected",
            "Itens afirmados pelo auditado selecionados para avaliacao.",
            quiet=args.quiet,
            identity=identity,
            total_itens=len(itens),
            itens=[item.codigo for item in itens],
            **base_log,
        )
        if not itens:
            result = {
                "status": "completed",
                "conclusoes": [],
                "skip_reason": "nenhum_item_afirmado",
            }
            gravar_registro_analise(
                checkpoint,
                {
                    "identity": identity,
                    "status": result["status"],
                    "auditado": analise.auditado,
                    "questao": questao_base,
                    "coluna_evidencia": analise.coluna_evidencia,
                    "evidencia": analise.nome_original_evidencia,
                    "provider": args.provider,
                    "model": args.model,
                    "result": result,
                    "error": "",
                    "finished_at": dt.datetime.now(dt.timezone.utc).isoformat(),
                },
            )
            registros[identity] = {"identity": identity, "status": result["status"]}
            total_processadas += 1
            total_concluidas += 1
            log_event(
                "analysis_skipped_no_items",
                "Analise concluida sem chamada ao provider porque nenhum item foi afirmado.",
                quiet=args.quiet,
                identity=identity,
                checkpoint=str(checkpoint),
                **base_log,
            )
            continue
        pacote = normalizar_evidencia(resolucao.caminho)
        log_event(
            "evidence_normalized",
            "Evidencia normalizada para envio ao provider.",
            quiet=args.quiet,
            level="warning" if pacote.erro else "info",
            identity=identity,
            tipo=pacote.tipo,
            documentos=len(pacote.documentos),
            inventario=len(pacote.inventario),
            error=pacote.erro,
            **base_log,
        )
        env_key = {"gemini": "GEMINI_API_KEY", "openrouter": "OPENROUTER_API_KEY"}.get(args.provider, "")
        api_key = os.environ.get(env_key, "") if env_key else ""
        with tempfile.TemporaryDirectory() as upload_tmp:
            arquivos_upload = arquivos_compativeis_upload(resolucao.caminho, upload_tmp)
            log_event(
                "upload_prepared",
                "Arquivos preparados para upload ao provider.",
                quiet=args.quiet,
                identity=identity,
                total_arquivos=len(arquivos_upload),
                arquivos=[Path(arquivo).name for arquivo in arquivos_upload],
                **base_log,
            )
            if args.provider in REMOTE_PROVIDERS and api_key:
                wait_seconds = rate_limiter.wait_seconds()
                if wait_seconds > 0:
                    log_event(
                        "rate_limit_wait",
                        "Aguardando limite de requests por minuto antes da chamada ao provider.",
                        quiet=args.quiet,
                        identity=identity,
                        provider=args.provider,
                        model=args.model,
                        rpm=args.rpm,
                        wait_seconds=round(wait_seconds, 3),
                        resposta_id=analise.resposta_id,
                        itens=[item.codigo for item in itens],
                        **base_log,
                    )
                rate_limiter.wait_and_mark(wait_seconds)
            log_event(
                "provider_started",
                "Chamada ao provider iniciada.",
                quiet=args.quiet,
                identity=identity,
                provider=args.provider,
                model=args.model,
                **base_log,
            )
            result = executar_provider(
                provider=args.provider,
                model=args.model,
                api_key=api_key,
                prompt=prompt.conteudo,
                auditado=analise.auditado,
                questao_base=questao_base,
                coluna_evidencia=analise.coluna_evidencia,
                itens_afirmados=itens,
                pacote={
                    "documentos": pacote.documentos,
                    "inventario": pacote.inventario,
                    "erro": pacote.erro,
                    "arquivos_upload": arquivos_upload,
                },
            )
        conclusoes = result.get("conclusoes") if isinstance(result, dict) else None
        log_event(
            "provider_finished",
            "Chamada ao provider finalizada.",
            quiet=args.quiet,
            level="error" if result.get("status") == "error" else "info",
            identity=identity,
            provider=args.provider,
            model=args.model,
            status=result.get("status"),
            conclusoes=len(conclusoes or []),
            error=result.get("error", ""),
            **base_log,
        )
        gravar_registro_analise(
            checkpoint,
            {
                "identity": identity,
                "status": result["status"],
                "auditado": analise.auditado,
                "questao": questao_base,
                "coluna_evidencia": analise.coluna_evidencia,
                "evidencia": analise.nome_original_evidencia,
                "provider": args.provider,
                "model": args.model,
                "result": result,
                "error": result.get("error", ""),
                "finished_at": dt.datetime.now(dt.timezone.utc).isoformat(),
            },
        )
        registros[identity] = {"identity": identity, "status": result["status"]}
        total_processadas += 1
        if result["status"] == "error":
            total_erros += 1
        else:
            total_concluidas += 1
        log_event(
            "analysis_recorded",
            "Resultado da analise gravado no checkpoint.",
            quiet=args.quiet,
            identity=identity,
            status=result["status"],
            checkpoint=str(checkpoint),
            **base_log,
        )
    relatorio = out_dir / "relatorio_conformidade.xlsx"
    linhas_relatorio = gerar_relatorio_conformidade(checkpoint, relatorio)
    log_event(
        "report_generated",
        "Relatorio de conformidade gerado.",
        quiet=args.quiet,
        relatorio=str(relatorio),
        linhas=linhas_relatorio,
    )
    log_event(
        "pipeline_finished",
        "Execucao do pipeline finalizada.",
        quiet=args.quiet,
        total_analises=len(analises),
        processadas=total_processadas,
        puladas=total_puladas,
        concluidas=total_concluidas,
        erros=total_erros,
        checkpoint=str(checkpoint),
        relatorio=str(relatorio),
    )
    return 0


def _linha_por_id(caminho_xlsx: str | Path, resposta_id: Any) -> dict[str, Any]:
    for linha in _rows_from_xlsx(Path(caminho_xlsx)):
        if linha.get("id") == resposta_id:
            return linha
    return {}
